import fitz
import camelot
import gc
import os
import shutil
import tempfile
import re
import unicodedata
from docx import Document

def extract_pdf_text_and_tables(pdf_path):
    """Extract text and tables safely from a PDF without locking files."""

    all_text = []

    # ---- TEXT EXTRACTION ----
    try:
        with fitz.open(pdf_path) as doc:
            for page in doc:
                all_text.append(page.get_text("text") + "\n")
    except Exception as e:
        print(f"Text extraction failed for {pdf_path}: {e}")

    # ---- TABLE EXTRACTION ----
    temp_dir = tempfile.mkdtemp()  # Use a private temp directory for Camelot
    try:
        tables = camelot.read_pdf(pdf_path, pages='all', flavor='stream', output_dir=temp_dir)
        for i, table in enumerate(tables):
            df = table.df
            table_text = df.to_string(index=False, header=True)
            all_text.append(f"\n[Table {i+1}]\n{table_text}\n")
    except Exception as e:
        print(f"Table extraction skipped for {pdf_path}: {e}")
    finally:
        # Clean up Camelot’s temp folder safely
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception as cleanup_err:
            print(f"Could not clean temp dir {temp_dir}: {cleanup_err}")

    # ---- MEMORY CLEANUP ----
    gc.collect()

    return "\n".join(all_text).strip()


def extract_docx_text_and_tables(docx_path):
    doc = Document(docx_path)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join([cell.text for cell in row.cells])
            full_text.append(row_text)

    # Join with newlines
    return "\n".join(full_text).strip()



def clean_text(text):
    """Remove or normalize weird unicode characters from PDF text."""
    # Normalize text to NFKC form (standardizes characters)
    text = unicodedata.normalize("NFKC", text)

    # Remove leftover non-printable characters
    text = re.sub(r"[^\x20-\x7E\n]+", "", text)

    # Collapse excessive whitespace
    text = re.sub(r"\s+", " ", text).strip()
    return text

